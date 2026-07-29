"""
doc_loader.py — Loaders for .docx and .txt, plus a format dispatcher
that routes to pdf_loader for PDFs.
"""

import logging
import os

from langchain_core.documents import Document

logger = logging.getLogger(__name__)

# Word has no fixed "pages" in the text — approximate a page break every
# N characters so chunks still carry a rough, human-meaningful locator.
DOCX_PAGE_CHAR_APPROX = int(os.environ.get("DOCX_PAGE_CHAR_APPROX", "1800"))


def load_txt(file_path: str) -> list[Document]:
    name = os.path.basename(file_path)
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read().strip()
    except Exception:
        logger.exception("Failed to read TXT file: %s", name)
        return []

    if not text:
        return []

    return [Document(page_content=text, metadata={"source": name, "page": 1})]


def load_docx(file_path: str) -> list[Document]:
    name = os.path.basename(file_path)

    try:
        import docx  # python-docx
    except ImportError:
        logger.error("python-docx is not installed; cannot read %s", name)
        return []

    try:
        document = docx.Document(file_path)
    except Exception:
        logger.exception("Failed to open DOCX file: %s", name)
        return []

    documents: list[Document] = []
    page_number = 1
    buffer: list[str] = []
    char_count = 0

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        buffer.append(text)
        char_count += len(text)

        if char_count >= DOCX_PAGE_CHAR_APPROX:
            documents.append(
                Document(
                    page_content="\n".join(buffer),
                    metadata={"source": name, "page": page_number},
                )
            )
            buffer = []
            char_count = 0
            page_number += 1

    if buffer:
        documents.append(
            Document(
                page_content="\n".join(buffer),
                metadata={"source": name, "page": page_number},
            )
        )

    # Tables often hold the most useful structured content (specs, resumes, etc.)
    for table in document.tables:
        rows_text = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows_text.append(" | ".join(cells))
        if rows_text:
            documents.append(
                Document(
                    page_content="\n".join(rows_text),
                    metadata={"source": name, "page": page_number},
                )
            )

    return documents


def load_document(file_path: str, original_filename: str) -> list[Document]:
    """Dispatch to the correct loader based on the *original* filename's extension,
    and stamp every returned Document with that original filename as its source
    (the on-disk path is a temp file with a random name)."""
    ext = os.path.splitext(original_filename)[1].lower()

    if ext == ".pdf":
        from pdf_loader import load_pdf
        documents = load_pdf(file_path)
    elif ext == ".docx":
        documents = load_docx(file_path)
    elif ext == ".txt":
        documents = load_txt(file_path)
    else:
        logger.warning("Unsupported file type for %s: %s", original_filename, ext)
        return []

    for doc in documents:
        doc.metadata["source"] = original_filename

    return documents
